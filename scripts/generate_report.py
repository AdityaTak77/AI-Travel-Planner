from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path
from typing import List

from PIL import Image, ImageDraw, ImageFont


def add_footer_page_numbers(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    # Insert current page number field
    fld_page = OxmlElement('w:fldSimple')
    fld_page.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
    run._r.append(fld_page)
    p.add_run(" of ")
    # Insert total pages field
    run2 = p.add_run("")
    fld_numpages = OxmlElement('w:fldSimple')
    fld_numpages.set(qn('w:instr'), 'NUMPAGES \\* MERGEFORMAT')
    run2._r.append(fld_numpages)


def add_title_page(document: Document, title: str, subtitle: str | None = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    if subtitle:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(12)
    # Date
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().strftime('%B %d, %Y'))


def add_toc(document: Document) -> None:
    document.add_heading('Contents', level=1)
    # Insert a TOC field (Word needs Update Field to render)
    p = document.add_paragraph()
    r = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    r._r.append(fld)


def add_paragraphs(document: Document, text: str) -> None:
    # Split on blank lines to make readable paragraphs
    for block in [b.strip() for b in text.split('\n\n') if b.strip()]:
        document.add_paragraph(block)


def _try_load_mono_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    # Try common Windows monospace fonts, then fall back
    candidates = [
        "consola.ttf",  # Consolas
        "Courier New.ttf",
        "C:\\Windows\\Fonts\\consola.ttf",
        "C:\\Windows\\Fonts\\cour.ttf",
        "C:\\Windows\\Fonts\\courbd.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def render_ascii_to_image(ascii_text: str, out_path: Path, padding: int = 20, font_size: int = 16) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    lines = ascii_text.splitlines() or [""]
    font = _try_load_mono_font(font_size)
    # Measure text
    dummy_img = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(dummy_img)
    max_w = 0
    line_h = 0
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        max_w = max(max_w, w)
        line_h = max(line_h, h)
    img_w = max_w + padding * 2
    img_h = line_h * len(lines) + padding * 2
    img = Image.new("RGB", (max(1, img_w), max(1, img_h)), "white")
    draw = ImageDraw.Draw(img)
    y = padding
    for line in lines:
        draw.text((padding, y), line, fill="black", font=font)
        y += line_h
    img.save(out_path)
    return out_path


def extract_ascii_diagrams_from_text(text: str) -> List[str]:
    # Capture fenced code blocks that contain box-drawing characters or typical diagram edges
    blocks: List[str] = []
    in_block = False
    cur: List[str] = []
    for line in text.splitlines():
        if line.strip().startswith("```"):
            if in_block:
                # closing
                block = "\n".join(cur).strip("\n")
                if any(ch in block for ch in ["┌", "┐", "└", "┘", "│", "─"]) or (
                    "+-" in block or "|" in block or "-" in block
                ):
                    blocks.append(block)
                cur = []
                in_block = False
            else:
                in_block = True
                cur = []
        elif in_block:
            cur.append(line.rstrip("\n"))
    # Also try to capture unfenced ASCII diagrams by scanning for box-drawing chars in runs
    lines = text.splitlines()
    buf: List[str] = []
    has_diagram = False
    for line in lines:
        if any(ch in line for ch in ["┌", "┐", "└", "┘", "│", "─"]):
            buf.append(line.rstrip("\n"))
            has_diagram = True
        else:
            if has_diagram:
                blocks.append("\n".join(buf).strip("\n"))
                buf = []
                has_diagram = False
    if has_diagram and buf:
        blocks.append("\n".join(buf).strip("\n"))
    return blocks


def add_workflow_diagrams(document: Document) -> None:
    document.add_heading('Workflow Diagrams', level=1)
    roots = [
        Path('README.md'),
        Path('MCP_ARCHITECTURE.txt'),
        Path('MCP_PROTOCOL.md'),
    ]
    output_dir = Path('docs/diagrams')
    count = 0
    for src in roots:
        if not src.exists():
            continue
        try:
            text = src.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            continue
        blocks = extract_ascii_diagrams_from_text(text)
        if not blocks:
            continue
        document.add_heading(f'{src.name}', level=2)
        for idx, block in enumerate(blocks, start=1):
            img_path = output_dir / f"{src.stem}_diagram_{idx}.png"
            try:
                render_ascii_to_image(block, img_path)
                document.add_paragraph(f'Diagram {idx}:')
                document.add_picture(str(img_path), width=Inches(6.5))
                count += 1
            except Exception:
                # Fallback to monospaced text if rendering fails
                p = document.add_paragraph()
                run = p.add_run('\n' + block + '\n')
                run.font.name = 'Courier New'
    if count == 0:
        document.add_paragraph('No diagrams detected to embed.')


def build_report() -> Document:
    doc = Document()

    add_footer_page_numbers(doc)
    add_title_page(doc, 'AI Travel Planner — Project Report', 'Multi-Agent MCP-Integrated System')
    doc.add_page_break()

    # Contents (TOC)
    add_toc(doc)
    doc.add_paragraph('Note: Page numbers will update after opening in Word: References > Update Table.')
    doc.add_paragraph('Required layout: Introduction (1-2), Problem Statement (3-4), Literature Review (5-10), Methodology (11-15), Technology Stack (16), Results (17-23), Conclusions (24), References, Appendix.')
    doc.add_page_break()

    # Workflow Diagrams Section (auto-extracted from project docs)
    add_workflow_diagrams(doc)
    doc.add_page_break()

    # Section: Introduction
    doc.add_heading('Introduction', level=1)
    intro = (
        "This report presents the design and implementation of an AI-powered Travel Planner that\n"
        "generates personalized itineraries using a multi-agent architecture, standardized tool\n"
        "integration via the Model Context Protocol (MCP), and structured data models with Pydantic.\n"
        "The system orchestrates research, itinerary generation, cost optimization, and monitoring,\n"
        "leveraging external services like Google Gemini, Groq, and DuckDuckGo.\n\n"
        "Key goals include: (1) reliable tool discovery and invocation, (2) traceable and auditable\n"
        "workflows, (3) reproducible structured outputs for itineraries, and (4) clear separation of\n"
        "concerns across agents, integrations, state, and logging."
    )
    add_paragraphs(doc, intro)

    # Section: Problem Statement
    doc.add_heading('Problem Statement', level=1)
    problem = (
        "Travel planning is time-consuming, fragmented across multiple information sources, and prone\n"
        "to inconsistencies. Users must research destinations, compare lodging, plan day-by-day\n"
        "activities, and budget across transport, accommodation, and experiences. Traditional tools\n"
        "lack: (a) standardized access to external services, (b) composable agent workflows, and\n"
        "(c) structured, auditable outputs.\n\n"
        "This project addresses these gaps by: (1) using MCP for type-safe tool integration, (2) a\n"
        "multi-agent approach (CrewAI and ADK agents) for planning and optimization, and (3) Pydantic\n"
        "models for consistent data exchange and storage."
    )
    add_paragraphs(doc, problem)

    # Section: Literature Review
    doc.add_heading('Literature Review', level=1)
    literature = (
        "Large Language Models (LLMs) have demonstrated strong capabilities for text generation and\n"
        "planning, but practical systems require tool-use to incorporate real-world data. Protocols\n"
        "like MCP formalize tool discovery, request/response schemas, and error handling, improving\n"
        "agent reliability. Structured modeling (e.g., Pydantic v2) enables schema validation and\n"
        "type safety at boundaries.\n\n"
        "Multi-agent coordination is an emerging pattern for decomposing complex tasks into expert\n"
        "roles (e.g., proposal generation vs. cost optimization). Observability—via correlation and\n"
        "trace IDs—helps diagnose failures and validate outcomes. Prior work on web search integration\n"
        "(DuckDuckGo), generative APIs (Groq, Gemini), and local utilities (calculators) shows that\n"
        "combining APIs with LLM reasoning leads to higher-quality, verifiable itineraries.\n\n"
        "This project builds on these ideas, providing: (1) an MCP client and adapters, (2) unified\n"
        "tool schemas for research/generation/search/calculation, and (3) a demo and tests showing\n"
        "discovery, invocation, and error handling."
    )
    add_paragraphs(doc, literature)

    # Section: Methodology
    doc.add_heading('Methodology', level=1)
    methodology = (
        "Architecture: The system follows a modular layout. User interaction flows through an\n"
        "interactive planner and workflows. The MCP client provides standardized discovery and\n"
        "invocation for four tools: Gemini research, Groq LLM, DuckDuckGo search, and a budget\n"
        "calculator. Adapters wrap each integration with async support, validations, and structured\n"
        "responses.\n\n"
        "Agents: A CrewAI-based agent assembles a destination prompt using MCP-provided research,\n"
        "calls Groq via MCP to generate an itinerary, then shares a proposal through the A2A protocol.\n"
        "An ADK-style optimizer validates cost and constraints, invoking the calculator tool as needed.\n\n"
        "Observability: Callbacks emit monitoring events with trace/correlation IDs. A JSON logger\n"
        "captures structured events for audits. Tests cover models, protocol envelope integrity, state\n"
        "store behavior, and integration flows.\n\n"
        "Data Models: Pydantic models define itineraries, events, and protocol envelopes, ensuring\n"
        "consistent serialization and validation across components."
    )
    add_paragraphs(doc, methodology)

    # Section: Technology Stack
    doc.add_heading('Technology Stack', level=1)
    tech = (
        "Core: Python 3.13, Pydantic v2, pytest, asyncio.\n\n"
        "MCP: Custom MCP client (tool registry, discovery, request/response envelopes) and adapters\n"
        "for four tools.\n\n"
        "Integrations:\n"
        "- Gemini 2.0 Flash (research)\n"
        "- Groq LLM (itinerary generation)\n"
        "- DuckDuckGo (web search)\n"
        "- Calculator (local budget math)\n\n"
        "Tooling: httpx, python-dotenv, ruff/black/mypy, pytest-asyncio, pytest-cov."
    )
    add_paragraphs(doc, tech)

    # Section: Results
    doc.add_heading('Results', level=1)
    results = (
        "MCP integration exposes four tools with validated schemas. A demo script demonstrates tool\n"
        "discovery, schema inspection, and sample invocations (e.g., calculator operations). The\n"
        "interactive planner successfully orchestrates MCP research and generation, producing structured\n"
        "itineraries and saving outputs to examples/. Monitoring events confirm agent progress and\n"
        "message exchange via the A2A protocol.\n\n"
        "Representative outcomes: (1) Successful discovery of 4 MCP tools; (2) Async invocation with\n"
        "robust coroutine handling; (3) Planner runs generating itineraries for sample destinations\n"
        "(e.g., Goa, Jaipur); (4) Budget calculations returning totals and per-day costs."
    )
    add_paragraphs(doc, results)

    # Section: Conclusions
    doc.add_heading('Conclusions', level=1)
    conclusions = (
        "The project delivers a practical, extensible AI Travel Planner that combines MCP-based tool\n"
        "integration, multi-agent workflows, and structured models. The approach improves reliability,\n"
        "observability, and maintainability compared to ad-hoc tool usage. Future work includes adding\n"
        "more data sources (e.g., booking APIs), strengthening optimization strategies, and expanding\n"
        "tests and benchmarks under real-world conditions."
    )
    add_paragraphs(doc, conclusions)

    # Section: References
    doc.add_heading('References', level=1)
    refs = [
        "[1] Model Context Protocol (MCP) – Project documentation",
        "[2] Pydantic v2 – https://docs.pydantic.dev/",
        "[3] Google Gemini API – https://ai.google.dev/",
        "[4] Groq API – https://groq.com/",
        "[5] DuckDuckGo Search – https://duckduckgo.com/",
        "[6] pytest – https://docs.pytest.org/",
        "[7] httpx – https://www.python-httpx.org/",
        "[8] python-docx – https://python-docx.readthedocs.io/",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # Section: Appendix
    doc.add_heading('Appendix', level=1)
    appendix = (
        "Included code and outputs are available in the repository. Key paths:\n\n"
        "- Source code: src/\n"
        "- MCP client and adapters: src/integrations/mcp_client.py, src/integrations/mcp_tool_adapter.py\n"
        "- Interactive planner: src/interactive_planner.py\n"
        "- Examples (generated itineraries): examples/\n"
        "- Tests: src/tests/\n"
        "- Documentation: README.md, MCP_PROTOCOL.md, MCP_IMPLEMENTATION_SUMMARY.md\n\n"
        "To re-run the planner: python -m src.interactive_planner\n"
        "To run the MCP demo: python examples/mcp_demo.py"
    )
    add_paragraphs(doc, appendix)

    return doc


if __name__ == '__main__':
    document = build_report()
    out_path = 'AI_Travel_Planner_Project_Report.docx'
    try:
        document.save(out_path)
        print(f'Report generated: {out_path}')
    except PermissionError:
        # Fallback if the file is open; save with a timestamped name
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        alt_path = f'AI_Travel_Planner_Project_Report_{ts}.docx'
        document.save(alt_path)
        print(f'Report in use; saved as: {alt_path}')
